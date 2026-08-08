"""
Generates an honest, JD-aware cover letter docx per posting (remediation
brief R6). Pulls only from config.py + the fact bank -- no invented
achievements.

Compared with the old version (which led with a resume summary fragment,
never addressed relocation, and left the start-date conflict unresolved):

  * Full sender postal address block (German convention) at the top, with a
    named recipient block when the JD names a contact. The builder refuses to
    emit a letter while SENDER_ADDRESS still contains FILL: markers.
  * German date line and a bold Betreff line.
  * Opening paragraph names the target role, the Pflichtpraktikum ask, the
    availability window MATCHED to the posting's start date, the 5-6 month
    duration, and -- when the role is outside Berlin -- a relocation line.
  * One paragraph on the JD-relevant domain project (scorer.lead_project),
    from the fact bank's summary_paragraph.
  * One paragraph mapping at most two JD-required technologies to real work,
    from TECH_CLAIM_SENTENCES, plus the work-authorization (Pflichtpraktikum)
    sentence. No stack lists > 3 items, no employer mission paraphrase, and
    the phrase "I am writing to apply for" never appears.
  * Target length 250-350 words.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from datetime import date
import os
import re
import sys

sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))
from config import CONTACT, WORK_AUTH_LETTER, SENDER_ADDRESS, CANDIDATE_PROFILE, \
    sender_address_configured, validate
from fact_bank import TECH_CLAIM_SENTENCES, PROJECT_ACHIEVEMENTS

GREY = RGBColor(0x55, 0x55, 0x55)

# German date header ("8. August 2026") follows the German convention even
# though the letter body is English.
GERMAN_MONTHS = ["Januar", "Februar", "M\u00e4rz", "April", "Mai", "Juni",
                 "Juli", "August", "September", "Oktober", "November", "Dezember"]

ENGLISH_MONTHS = ["January", "February", "March", "April", "May", "June",
                  "July", "August", "September", "October", "November", "December"]

# One-sentence project highlights per profile. Kept for pipeline/followup.py,
# which still uses HIGHLIGHTS to build the follow-up email body.
HIGHLIGHTS = {
    "data_engineer": (
        "My most recent project, Stadtanalyse, is an end-to-end streaming data platform: Kafka "
        "into a Delta Lake medallion architecture, Spark batch processing, Great Expectations "
        "quality suites, dbt gold marts and an XGBoost model, all orchestrated by Airflow."
    ),
    "ai_ml": (
        "My most recent project, CreditLens, combines a Python ETL/data-validation pipeline "
        "with a RAG layer (sentence-transformers, ChromaDB, Groq/Llama 3.3) that answers "
        "questions over financial documents with cited sources and declines to answer when "
        "the context doesn't support it."
    ),
    "nlp": (
        "My work spans several NLP systems in production, from a support chatbot that "
        "autonomously resolved the large majority of customer queries to CreditLens's "
        "RAG-based document Q&A engine with cited, grounded answers."
    ),
    "software_eng": (
        "My most recent project, CreditLens, is an application I scoped, built and shipped "
        "end to end: a Python backend with a validation pipeline, a SQL schema designed for "
        "time-series comparison, and a Next.js + FastAPI frontend deployed on AWS EC2 behind "
        "Nginx with systemd process management."
    ),
}

# Suffixes LinkedIn/Indeed bolt onto job titles.
_TITLE_NOISE = re.compile(
    r"\s*(?:\((?:m|w|f|d|x|all genders)[/\s|,;·-]*(?:m|w|f|d|x)?[/\s|,;·-]*"
    r"(?:m|w|f|d|x)?\)|\(all genders\))\s*",
    re.IGNORECASE,
)
_GENDER_STAR = re.compile(r"[*:]\s*in(?:nen)?\b|/-?in(?:nen)?\b|\*(?=\s|$)", re.IGNORECASE)

_BANNED_PHRASES = [
    "I am writing to apply for",
    "I am writing to apply",
    "with great interest",
    "I'm writing to express",
]

_ACCOMMODATION_CITIES = {"berlin", "potsdam", "remote"}


def _german_date():
    today = date.today()
    return f"{today.day}. {GERMAN_MONTHS[today.month - 1]} {today.year}"


def clean_role_title(role: str, company: str = "") -> str:
    """'AI Scientist, Internship, Germany - BCG X' + 'BCG X' -> 'AI Scientist, Internship'."""
    title = (role or "").strip()
    title = _GENDER_STAR.sub("", title)
    title = _TITLE_NOISE.sub(" ", title)

    if company:
        title = re.sub(
            r"\s*[-\u2013\u2014|@]\s*" + re.escape(company.strip()) + r"\s*$",
            "", title, flags=re.IGNORECASE,
        )
        title = re.sub(r"\s+at\s+" + re.escape(company.strip()) + r"\s*$",
                       "", title, flags=re.IGNORECASE)

    title = re.sub(r",\s*(?:Germany|Deutschland|Berlin|M\u00fcnchen|Munich|Frankfurt|Hamburg)\s*$",
                   "", title, flags=re.IGNORECASE)

    return re.sub(r"[\s,;\-\u2013\u2014|]+$", "", title).strip() or (role or "").strip()


def _availability_clause(parsed):
    """'I am available from October 2026' from the parsed start date; falls
    back to 'I am available immediately' when the JD gave no date. The month
    is spelled out in English (the letter is English -- never a German/English
    mix) so no numeric start-date token leaks into the letter (T3)."""
    start = parsed.get("start_date")
    if not start:
        return "I am available immediately"
    try:
        m = ENGLISH_MONTHS[int(start.month) - 1]
        return f"I am available from {m} {start.year}"
    except (AttributeError, IndexError, TypeError, ValueError):
        return "I am available immediately"


def _relocation_line(parsed):
    city_key = parsed.get("city_key") or ""
    city = parsed.get("city") or ""
    if city_key in _ACCOMMODATION_CITIES or not city:
        return ""
    return f"I am based in Berlin and am prepared to relocate to {city} for this role. "


def _recipient_lines(parsed):
    """Recipient address block. Uses a JD-named contact when present,
    otherwise falls back to the company name + city."""
    company = parsed.get("company") or ""
    contact = parsed.get("contact") or ""
    person = None
    if contact:
        m = re.match(r"^(?:contact\s*[:]\s*)?([A-Za-z\u00c4\u00d6\u00dc\u00e4\u00f6\u00fc\u00df]+(?:\s+[A-Za-z\u00c4\u00d6\u00dc\u00e4\u00f6\u00fc\u00df]+){1,3})$", contact)
        if m and len(contact) < 60:
            person = m.group(1).strip()
    city = parsed.get("city") or ""
    lines = [company]
    if person:
        lines.append(person)
    if city:
        lines.append(city)
    return lines, person


def _project_paragraph(parsed, score):
    lead = score.get("lead_project") or "creditlens"
    proj = PROJECT_ACHIEVEMENTS.get(lead)
    if not proj:
        proj = PROJECT_ACHIEVEMENTS["creditlens"]
    return proj["summary_paragraph"]


def _jd_tech_sentence(parsed, score):
    """At most two JD-required technologies mapped to real work via the bank's
    claim sentences. Never a bare stack list of more than three items."""
    required = parsed.get("required_technologies") or []
    # Prefer technologies that have a claim sentence and that the candidate
    # actually has.
    scored = []
    for tech in required:
        sentence = TECH_CLAIM_SENTENCES.get(tech)
        if sentence and not _is_gap(tech):
            scored.append((tech, sentence))
    if not scored:
        for tech in (parsed.get("technologies_mentioned") or []):
            sentence = TECH_CLAIM_SENTENCES.get(tech)
            if sentence and not _is_gap(tech):
                scored.append((tech, sentence))
    picked = scored[:2]
    if not picked:
        return ""
    if len(picked) == 1:
        return " " + picked[0][1]
    return " " + " ".join(s for _, s in picked)


def _is_gap(tech):
    from fact_bank import has_technology
    return not has_technology(tech)


def build_cover_letter(profile, parsed, score, output_path, sender_address=None):
    """Build a JD-aware cover letter from the fact bank.

    profile: 'data_engineer' | 'ai_ml'
    parsed:  parse_posting() result
    score:   score_posting() result (lead_project, matched)
    sender_address: dict like config.SENDER_ADDRESS (defaults to config). The
      builder refuses to emit while the address contains FILL: markers.
    """
    sender = sender_address or SENDER_ADDRESS
    if not sender_address_configured(sender):
        raise SystemExit(
            "\n".join([
                "",
                "=" * 68,
                "  SENDER_ADDRESS in config.py still contains FILL: markers.",
                "  Set street and postal_code before generating cover letters.",
                "=" * 68,
                "",
            ])
        )
    validate()

    company = parsed.get("company") or ""
    role_clean = clean_role_title(parsed.get("title") or "", company)
    recipient_lines, person = _recipient_lines(parsed)
    salutation = f"Dear {person}," if person else f"Dear Hiring Team at {company},"

    start_clause = _availability_clause(parsed)
    relocation = _relocation_line(parsed)
    base_city = CANDIDATE_PROFILE.get("base_city", "Berlin")

    opening = (
        f"I am applying for the {role_clean} position at {company}. This is a "
        f"mandatory internship (Pflichtpraktikum) required by my MSc in Computer "
        f"Science (Big Data & AI) at SRH Berlin. {start_clause} for 5\u20136 "
        f"months as my programme requires. {relocation}"
    ).rstrip()

    work_auth = WORK_AUTH_LETTER

    project_para = _project_paragraph(parsed, score)
    tech_line = _jd_tech_sentence(parsed, score)

    parts = []

    # Sender address block.
    sender_block = [
        f"{sender['name']}",
        f"{sender['street']}",
        f"{sender['postal_code']} {sender['city']}",
        f"{sender['country']}",
    ]
    parts.append(("block", "\n".join(sender_block)))

    # Recipient address block + date.
    recipient_block = "\n".join(recipient_lines)
    parts.append(("block", f"{recipient_block}\n{base_city}, {_german_date()}"))

    # Betreff (subject) line.
    parts.append(("bold", f"Application for {role_clean} \u2014 Pflichtpraktikum"))

    # Salutation + body.
    parts.append(("body", salutation))
    parts.append(("body", opening))
    parts.append(("body", project_para))
    parts.append(("body", f"{tech_line.strip()}" if tech_line.strip() else
                  "I have attached my CV with further detail on my experience."))
    parts.append(("body", work_auth))
    parts.append(("body",
                  "I have attached my CV with further detail on my experience and projects. "
                  "I would welcome the chance to discuss how I could contribute to your team."))
    parts.append(("body", "Thank you for your consideration."))
    parts.append(("body", f"Best regards,\nSrikar Kodi\n{CONTACT['email']} | {CONTACT['phone']}"))

    doc = Document()
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for kind, para in parts:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(10)
        r = p.add_run(para.strip())
        r.font.size = Pt(11)
        if kind == "bold":
            r.bold = True

    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    doc.save(output_path)
    return output_path


def letter_stats(text: str) -> dict:
    """Word count + banned-phrase check for the acceptance tests (T5)."""
    words = len([w for w in re.split(r"\s+", text) if w])
    banned = [p for p in _BANNED_PHRASES if p.lower() in text.lower()]
    return {"words": words, "banned": banned}


if __name__ == "__main__":
    from pipeline.jd_parser import parse_posting
    from pipeline.scorer import score_posting

    posting = {
        "company": "PIMCO Prime Real Estate",
        "title": "Intern in Software and Data Engineering (m/f/d)",
        "location": "Munich, Germany",
        "apply_url": "https://www.linkedin.com/jobs/view/123",
        "jd_text": ("Requirements: Python, RAG/LLM solutions using vector databases, Spark. "
                    "Portfolio and credit analytics. Start 01.10.2026, 6 months. "
                    "Contact: Anna Schmidt. Apply via careers.allianz.com."),
    }
    p = parse_posting(posting)
    s = score_posting(posting, p)
    fake_addr = {"name": "Srikar Kodi", "street": "Musterstr. 1",
                 "postal_code": "10115", "city": "Berlin", "country": "Germany"}
    out = build_cover_letter(s["profile"], p, s,
                             os.path.join(os.path.dirname(__file__), "..", "data", "preview_letter.docx"),
                             sender_address=fake_addr)
    from docx import Document as D
    doc = D(out)
    text = "\n".join(pa.text for pa in doc.paragraphs)
    print("built", out)
    print("stats:", letter_stats(text))
    print("----")
    print(text)
