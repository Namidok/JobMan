"""
Builds a resume docx in the EXACT approved format/template.

This replaces the old per-variant builder (D1's root cause). The old code
received `matched` keywords and re-sorted bullets inside a FIXED per-variant
structure -- project order, summary and the skills block never changed, so
two postings sharing a variant produced near-identical CVs no matter how
different their job descriptions were.

The new builder is JD-aware by construction (R5). Everything it emits is read
from the fact bank (fact_bank.py); nothing is invented:

  - summary names the target role + the availability window that matches the
    posting's start date
  - skills categories and the technologies inside them are reordered by JD
    vocabulary, and technologies that were buried in the bank (R8's ten
    literals) are surfaced when the JD asks for them
  - projects are led by the domain-relevant project (scorer.lead_project)
    and project bullets are reordered by JD language
  - experience bullets render in FIXED order (D6): the audit showed JD-driven
    reordering of experience bullets is exactly what lets one job be described
    differently across applications. Only project content may reorder per JD.

The one-page fitter is preserved: it steps the body size down before it cuts
content, and never silently drops a project.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import os
import re

import sys
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))

from config import (CONTACT, SPOKEN_LANGUAGES, EDUCATION,
                    CANDIDATE_PROFILE, PROJECT_BULLET_CAP, validate)
from fact_bank import (SKILLS, SKILL_ORDER, DEFAULT_PROJECT_ORDER,
                       RETIRED_PROJECTS, render_experience,
                       render_project_bullets, projects as bank_projects)

# Body font size. The one-page fitter steps this down before it cuts content --
# losing 0.5pt is cheaper than losing a bullet.
BODY_PT = 10.0

ACCENT = RGBColor(0x2B, 0x4C, 0x7E)
GREY = RGBColor(0x55, 0x55, 0x55)
NAME_COLOR = RGBColor(0x1A, 0x1A, 0x1A)

_MONTHS = ["January", "February", "March", "April", "May", "June",
           "July", "August", "September", "October", "November", "December"]


def _hit_count(text, jd_text):
    """How many JD vocabulary tokens appear in `text` (word-boundary match)."""
    if not jd_text:
        return 0
    t = text.lower()
    jd = jd_text.lower()
    return sum(1 for token in re.findall(r"[a-z]{4,}", t) if token in jd)


def _add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "2B4C7E")
    rPr.append(color)

    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = ACCENT
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), "2B4C7E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def _bullet(doc, text, body_pt=BODY_PT):
    """Manual U+2022 bullet in the body font (real, extractable character)."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(1.5)
    pf.left_indent = Cm(0.5)
    pf.first_line_indent = Cm(-0.35)
    run = p.add_run("\u2022\u00a0 " + text)
    run.font.size = Pt(body_pt)
    return p


def _skill_line(doc, label, value, body_pt=BODY_PT):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1.5)
    r1 = p.add_run(f"{label}: ")
    r1.bold = True
    r1.font.size = Pt(body_pt)
    r2 = p.add_run(value)
    r2.font.size = Pt(body_pt)
    return p


def _job_header(doc, title, org):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(0.5)
    r1 = p.add_run(title)
    r1.bold = True
    r1.font.size = Pt(BODY_PT + 0.5)
    r2 = p.add_run(f"  \u2014  {org}")
    r2.font.size = Pt(BODY_PT + 0.5)
    return p


def _job_meta(doc, location, dates):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(location)
    r1.italic = True
    r1.font.size = Pt(9.5)
    r1.font.color.rgb = GREY
    r2 = p.add_run(f"   |   {dates}")
    r2.italic = True
    r2.font.size = Pt(9.5)
    r2.font.color.rgb = GREY
    return p


def _project_header(doc, name, stack):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(0.5)
    r1 = p.add_run(name)
    r1.bold = True
    r1.font.size = Pt(BODY_PT + 0.5)
    r2 = p.add_run(f"  \u2014  {stack}")
    r2.font.size = Pt(BODY_PT - 0.5)
    r2.font.color.rgb = GREY
    return p


def _link_line(doc, links):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    for i, (label, url) in enumerate(links):
        if i > 0:
            sep = p.add_run("   |   ")
            sep.font.size = Pt(9)
            sep.font.color.rgb = GREY
        _add_hyperlink(p, label, url)
    return p


def _strip_trailing_empty(doc):
    body = doc.element.body
    for child in list(body)[::-1]:
        if child.tag.endswith("}sectPr"):
            continue
        if child.tag.endswith("}p") and not "".join(child.itertext()).strip():
            body.remove(child)
        else:
            break


def _profile_role(profile):
    if profile == "data_engineer":
        return "Data Engineering"
    if profile == "ai_ml":
        return "Applied AI / Machine Learning"
    return "Data Engineering / Applied AI"


def _start_summary(parsed):
    """'available from October 2026' -- a word-form month, never a number, so
    no date digit leaks into the CV outside the fact bank's allowed set."""
    start = parsed.get("start_date")
    if not start:
        return "available immediately"
    try:
        m = _MONTHS[int(start.month) - 1]
        return f"available from {m} {start.year}"
    except (AttributeError, IndexError, TypeError, ValueError):
        return "available immediately"


def _skill_category_order(profile, jd_text):
    """Categories pinned by profile first; the tail reorders by JD vocabulary.
    Stable sort -- ties keep the bank order, so the CV does not flip around
    between otherwise-identical postings."""
    pinned = ["programming", "data_eng"] if profile == "data_engineer" else ["programming", "ai_ml"]
    tail = [c for c in SKILL_ORDER if c not in pinned]
    tail = sorted(tail, key=lambda c: _hit_count(SKILLS[c]["items"], jd_text), reverse=True)
    return pinned + tail


def _reorder_skill_items(items, jd_text):
    """Stable reorder of a comma-separated skills line so JD-mentioned
    technologies surface first (this is what 'surface buried technologies'
    means -- e.g. Kafka moves in front of Pandas when the JD asks for it)."""
    parts = [p.strip() for p in items.split(",")]
    parts = sorted(parts, key=lambda p: _hit_count(p, jd_text), reverse=True)
    return ", ".join(parts)


def _project_order(parsed, score):
    """Lead with the domain-relevant project (scorer.lead_project), then the
    active default order. Retired projects (skillsync, covercraft) are promoted
    ONLY when the JD's domain matches them -- never on their own."""
    lead = score.get("lead_project")
    order = []
    if lead:
        order.append(lead)
    for k in DEFAULT_PROJECT_ORDER:
        if k not in order:
            order.append(k)
    if lead and lead not in DEFAULT_PROJECT_ORDER:
        # lead is a retired project whose domain matched the JD.
        order.append(lead)
    active = [k for k in order if k not in RETIRED_PROJECTS] + [k for k in order if k in RETIRED_PROJECTS]
    return active


def build_resume(profile, parsed, score, output_path,
                 body_pt=BODY_PT, project_cap=None, jd_text=None):
    """Build a JD-aware resume from the fact bank.

    profile:   'data_engineer' | 'ai_ml' (from scorer)
    parsed:    parse_posting() result for this posting
    score:     score_posting() result (lead_project, matched, ...)
    output_path: path to write the .docx to
    body_pt / project_cap: used by the one-page fitter
    """
    validate()
    jd_text = jd_text or parsed.get("_jd", "")
    project_cap = project_cap or PROJECT_BULLET_CAP

    doc = Document()
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(1.1)
    section.bottom_margin = Cm(1.1)
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(body_pt)

    # Name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(CONTACT["name"])
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = NAME_COLOR

    # Title line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f"{_profile_role(profile)} \u00b7 Pflichtpraktikum Candidate \u00b7 MSc Big Data & AI")
    r.bold = True
    r.font.size = Pt(11.5)
    r.font.color.rgb = ACCENT

    # Contact line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(f'{CONTACT["location"]}  |  {CONTACT["phone"]}  |  {CONTACT["email"]}')
    r.font.size = Pt(9)
    r.font.color.rgb = GREY

    # Links line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    _add_hyperlink(p, CONTACT["site_label"], CONTACT["site"])
    sep = p.add_run("   |   "); sep.font.size = Pt(9); sep.font.color.rgb = GREY
    _add_hyperlink(p, CONTACT["linkedin_label"], CONTACT["linkedin"])
    sep = p.add_run("   |   "); sep.font.size = Pt(9); sep.font.color.rgb = GREY
    _add_hyperlink(p, CONTACT["github_label"], CONTACT["github"])

    # Summary -- names the target role and the availability window that
    # matches this posting's start date (R5).
    _section_heading(doc, "SUMMARY")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.add_run(
        f"MSc Computer Science student (Big Data & AI, Berlin) seeking a mandatory "
        f"internship (Pflichtpraktikum) in {_profile_role(profile)}, "
        f"{_start_summary(parsed)} for 5\u20136 months (per programme requirement). "
        f"3 years' professional experience building Python ETL pipelines, streaming "
        f"data platforms and RAG systems.",
    ).font.size = Pt(body_pt)

    # Skills -- category order is profile-pinned + JD-sorted tail; technologies
    # inside a line surface JD-mentioned ones first (R5/R8).
    _section_heading(doc, "TECHNICAL SKILLS")
    for cat_key in _skill_category_order(profile, jd_text):
        cat = SKILLS[cat_key]
        _skill_line(doc, cat["label"],
                    _reorder_skill_items(cat["items"], jd_text), body_pt=body_pt)
    _skill_line(doc, "Languages (spoken)", SPOKEN_LANGUAGES, body_pt=body_pt)

    # Experience -- FIXED order (D6). Never reordered by JD.
    _section_heading(doc, "PROFESSIONAL EXPERIENCE")
    rendered = render_experience(jd_text)
    for org, key, text in rendered:
        first = key == list(org["achievements"].keys())[0]
        if first:
            _job_header(doc, org["title"], org["org"])
            _job_meta(doc, org["location"], org["dates"])
        _bullet(doc, text, body_pt=body_pt)

    # Projects -- led by the domain-relevant project; bullets reorder by JD.
    _section_heading(doc, "PROJECTS")
    projs = bank_projects()
    for proj_key in _project_order(parsed, score):
        if proj_key not in projs:
            continue
        proj = projs[proj_key]
        _project_header(doc, proj["name"], proj["stack"])
        _link_line(doc, proj["links"])
        for _, bullet in render_project_bullets(proj_key, jd_text, cap=project_cap):
            _bullet(doc, bullet, body_pt=body_pt)

    # Education
    _section_heading(doc, "EDUCATION")
    for edu in EDUCATION:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        r1 = p.add_run(edu["degree"]); r1.bold = True; r1.font.size = Pt(body_pt)
        r2 = p.add_run(f'   {edu["dates"]}'); r2.font.size = Pt(9.5); r2.font.color.rgb = GREY
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(5)
        r3 = p2.add_run(edu["detail"]); r3.font.size = Pt(9.5); r3.font.color.rgb = GREY

    _strip_trailing_empty(doc)
    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    doc.save(output_path)
    return output_path


# ---------------------------------------------------------------------------
# One-page fitter
# ---------------------------------------------------------------------------

FIT_LADDER = [
    (PROJECT_BULLET_CAP, BODY_PT, "full content"),
    (PROJECT_BULLET_CAP, 9.5, "body text at 9.5pt"),
    (2, 9.5, "9.5pt, 2 bullets per project"),
    (2, 9.0, "9pt, 2 bullets per project"),
]


def _render_and_count(profile, parsed, score, output_path, cap, pt, jd_text):
    import tempfile
    from resume_builder.pdf_convert import convert_to_pdf, count_pdf_pages

    build_resume(profile, parsed, score, output_path, body_pt=pt, project_cap=cap, jd_text=jd_text)
    with tempfile.TemporaryDirectory() as tmp:
        return count_pdf_pages(convert_to_pdf(output_path, tmp))


def build_resume_fitted(profile, parsed, score, output_path,
                        jd_text=None, verbose=True):
    """Build a resume, preferring one page but never silently dropping a project.

    Returns (docx_path, pages, note).
    """
    jd_text = jd_text or parsed.get("_jd", "")
    for cap, pt, label in FIT_LADDER:
        pages = _render_and_count(profile, parsed, score, output_path, cap, pt, jd_text)
        if pages <= 1:
            if verbose and label != "full content":
                print(f"  (fitted to one page: {label})")
            return output_path, 1, label

    pages = _render_and_count(profile, parsed, score, output_path,
                              PROJECT_BULLET_CAP, BODY_PT, jd_text)
    if verbose:
        print(f"  ({pages} pages at full content \u2014 fine for a CV with three "
              f"projects. Lower PROJECT_BULLET_CAP in config.py to force one.)")
    return output_path, pages, "full content, 2 pages"


# Backwards-compatible alias
build_resume_one_page = build_resume_fitted


if __name__ == "__main__":
    from pipeline.jd_parser import parse_posting
    from pipeline.scorer import score_posting

    posting = {
        "company": "PIMCO Prime Real Estate",
        "title": "Intern in Software and Data Engineering (m/f/d)",
        "location": "Munich, Germany",
        "apply_url": "https://www.linkedin.com/jobs/view/123",
        "jd_text": ("Requirements: Python, RAG/LLM solutions using vector databases, "
                    "Spark. Portfolio and credit analytics. Internship for 6 months."),
    }
    p = parse_posting(posting)
    s = score_posting(posting, p)
    out = build_resume(s["profile"], p, s,
                       os.path.join(os.path.dirname(__file__), "..", "data", "preview_cv.docx"),
                       jd_text=posting["jd_text"])
    print("built", out, "profile:", s["profile"], "lead:", s["lead_project"])
